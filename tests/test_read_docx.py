"""read_docx — .docx supplement reader (tables + prose). Fixtures are built in-test with
python-docx so there's no dependency on gitignored data/raw. Covers the table summary, caption
association, specific-table preview (reusing read_table's header-skip/row_filter/columns/byte-cap
via _render_matrix), underline reveal, prose paging, and out-of-range handling."""
import json
import pathlib

import docx
import pytest

import agent_core


def _doc(path, captions_and_rows, prose=None):
    """captions_and_rows: list of (caption, rows[][]) -> paragraph caption then a table."""
    d = docx.Document()
    for cap, rows in captions_and_rows:
        if cap:
            d.add_paragraph(cap)
        t = d.add_table(rows=len(rows), cols=len(rows[0]))
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                t.rows[r].cells[c].text = str(val)
    for para in (prose or []):
        d.add_paragraph(para)
    d.save(str(path))
    return str(path)


# ---- summary + caption association ----

def test_summary_lists_tables_dims_and_captions(tmp_path):
    p = _doc(tmp_path / "s.docx", [
        ("Supplementary Table S6. New neoantigen mutations in recurrent tumor",
         [["Gene", "Mutation"], ["TP53", "R175H"], ["KRAS", "G12D"]]),
        ("Supplementary Table S1. Baseline characteristics", [["Age", "Sex"], ["61", "M"]]),
    ])
    out = agent_core.read_docx_from(p)  # no table_index -> summary
    assert "2 table(s)" in out
    assert "S6" in out and "S1" in out and "table_index" in out
    assert "3x2" in out and "2x2" in out  # dims rows×cols


def test_no_tables_doc_points_to_prose(tmp_path):
    p = _doc(tmp_path / "n.docx", [], prose=["Supplementary methods.", "Patients were…"])
    out = agent_core.read_docx_from(p)
    assert "0 tables" in out and "text_offset" in out


# ---- specific table ----

def test_read_specific_table_preview(tmp_path):
    p = _doc(tmp_path / "t.docx", [
        ("Table S6.", [["Gene", "Mutation"], ["TP53", "R175H"], ["KRAS", "G12D"]])])
    out = agent_core.read_docx_from(p, table_index=0)
    assert "showing" in out and "table[0]" in out
    body = json.loads(out[out.index("\n", out.index(":\n")) + 1:]) if False else json.loads(out.split(":\n", 1)[1])
    assert ["Gene", "Mutation"] in body and ["TP53", "R175H"] in body


def test_table_index_out_of_range(tmp_path):
    p = _doc(tmp_path / "t.docx", [("c", [["a"], ["b"]])])
    out = agent_core.read_docx_from(p, table_index=5)
    assert "out of range" in out


def test_columns_projection_with_title_row_skip(tmp_path):
    # first row is a single-cell title -> _header_index should skip it and use row 1 as header
    p = _doc(tmp_path / "t.docx", [
        ("cap", [["Table S6. neoantigens", ""], ["Gene", "Mutation"],
                 ["TP53", "R175H"], ["KRAS", "G12D"]])])
    out = agent_core.read_docx_from(p, table_index=0, columns=["Gene"])
    assert "data rows" in out
    body = json.loads(out.split(":\n", 1)[1])
    assert {"Gene": "TP53"} in body and {"Gene": "KRAS"} in body


def test_row_filter(tmp_path):
    p = _doc(tmp_path / "t.docx", [
        ("cap", [["Gene", "Mutation"], ["TP53", "R175H"], ["KRAS", "G12D"]])])
    out = agent_core.read_docx_from(p, table_index=0, row_filter={"col": "Gene", "equals": "KRAS"})
    body = json.loads(out.split(":\n", 1)[1])
    assert len(body) == 1 and body[0]["Mutation"] == "G12D"


# ---- underline reveal (minimal-epitope marking) ----

def test_underline_reveal(tmp_path):
    d = docx.Document()
    d.add_paragraph("Table S2.")
    t = d.add_table(rows=2, cols=1)
    t.rows[0].cells[0].text = "Peptide"
    cell = t.rows[1].cells[0]
    cell.paragraphs[0].text = ""
    r1 = cell.paragraphs[0].add_run("QLAS")
    r2 = cell.paragraphs[0].add_run("STYTAYIV"); r2.font.underline = True
    r3 = cell.paragraphs[0].add_run("GYV")
    p = str(tmp_path / "u.docx"); d.save(p)
    out = agent_core.read_docx_from(p, table_index=0, underline=True)
    assert "<u>STYTAYIV</u>" in out


# ---- prose paging ----

def test_prose_paging(tmp_path):
    long = "X" * 100
    p = _doc(tmp_path / "p.docx", [], prose=[long, long, long])  # ~302 chars joined
    win = agent_core.read_docx_from(p, text_offset=0, max_chars=120)
    assert "docx prose 0-120" in win and "text_offset=120" in win
    rest = agent_core.read_docx_from(p, text_offset=120, max_chars=1000)
    assert "120-" in rest and "text_offset=" not in rest.split("chars:\n", 1)[1]  # no further page


def _pivot_doc(path, header, columns):
    """Build the 27274999 Supp Table 1 shape: ONE header row + ONE data row whose cells each hold
    a whole column joined by newlines (the merged/pivoted layout that silently lost rows)."""
    d = docx.Document()
    t = d.add_table(rows=2, cols=len(header))
    for c, h in enumerate(header):
        t.rows[0].cells[c].text = h
    for c, col in enumerate(columns):
        t.rows[1].cells[c].text = "\n".join(col)
    d.save(str(path))
    return str(path)


# 31-peptide panel verbatim from PMID 27274999 Supplementary Table 1 (the real regression).
_PANEL_NAMES = ["CypB-129", "Lck-246", "Lck-422", "MAP-432", "WHSC2-103", "HNRPL-501", "UBE-43",
                "UBE-85", "WHSC2-141", "HNRPL-140", "SART3-302", "SART3-309", "SART2-93", "SART3-109",
                "Lck-208", "PAP-213", "PSA-248", "EGFR-800", "MRP3-503", "MRP3-1293", "SART2-161",
                "Lck-486", "Lck-488", "PSMA-624", "EZH2-735", "PTHrP-102", "SART3-511", "SART3-734",
                "Lck-90", "Lck-449", "PAP-248"]
_PANEL_SEQS = ["KLKHYGPGWV", "KLVERLGAA", "DVWSFGILL", "DLLSHAFFA", "ASLDSDPWV", "NVLHFFNAPL",
               "RLQEWCSVI", "LIADFLSGL", "ILGELREKV", "ALVEFEDVL", "LLQAEAPRL", "RLAEYQAYI",
               "DYSARWNEI", "VYDYNCHVDL", "HYTNASDGL", "LYCESVHNF", "HYRKWIKDTI", "DYVREHKDNI",
               "LYAWEPSFL", "NYSVRYRPGL", "AYDFLYNYL", "TFDYLRSVL", "DYLRSVLEDF", "TYSVSFDSL",
               "KYVGIEREM", "RYLTQETNKV", "WLEYYNLER", "QIRPIFSNR", "ILEQSGEWWK", "VIQNLERGYR",
               "GIHKQKEKSR"]


def test_unpivot_recovers_all_rows_from_list_cells(tmp_path):
    """A pivoted list-cell table (1 data row, all values newline-joined) must un-pivot to one row
    per value — NOT lose the tail to the per-cell clip (the 27274999 23/31 regression). One header
    cell WRAPS across lines ("Amino acid\\nsequence") with MIXED counts across the row, exactly like the
    real Supp Table 1 — the guard must key on the DATA row's uniform K, not bail on any header newline
    (the original bug bailed on ANY header newline and lost all 8 tail rows)."""
    p = _pivot_doc(tmp_path / "panel.docx", ["Peptide name", "Amino acid\nsequence"],
                   [_PANEL_NAMES, _PANEL_SEQS])
    out = agent_core.read_docx_from(p, table_index=0)
    assert "un-pivoted" in out
    for seq in _PANEL_SEQS:                      # every sequence survives, incl. the 8 once dropped
        assert seq in out, f"{seq} lost"
    proj = agent_core.read_docx_from(p, table_index=0, columns=["Amino acid\nsequence"])
    assert proj.count("WLEYYNLER") == 1 and proj.count("GIHKQKEKSR") == 1


def test_summary_flags_list_cell_pivot(tmp_path):
    p = _pivot_doc(tmp_path / "panel.docx", ["Name", "Seq"], [_PANEL_NAMES, _PANEL_SEQS])
    out = agent_core.read_docx_from(p)           # summary
    assert "un-pivoted" in out and "31 rows" in out


def test_normal_table_is_not_unpivoted(tmp_path):
    """A genuine 2-row table (header + one record) where one cell merely wraps to two lines must
    NOT be reshaped — the strict guard requires a CONSISTENT K>=2 across ALL cells."""
    d = docx.Document()
    t = d.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Gene"; t.rows[0].cells[1].text = "Note"
    t.rows[1].cells[0].text = "TP53"; t.rows[1].cells[1].text = "line one\nline two"
    p = str(tmp_path / "normal.docx"); d.save(p)
    out = agent_core.read_docx_from(p, table_index=0)
    assert "un-pivoted" not in out and "TP53" in out


def test_ambiguous_uniform_header_is_not_unpivoted(tmp_path):
    """When BOTH rows are uniform K>=2 lists (header can't be told from data), the reshape is
    ambiguous — bail rather than risk silently corrupting a normal table (conservative guard)."""
    p = _pivot_doc(tmp_path / "amb.docx", ["Col\nA", "Col\nB"],   # header: every cell wraps to 2
                   [["x", "y"], ["p", "q"]])                       # data: every cell also 2 -> ambiguous
    out = agent_core.read_docx_from(p, table_index=0)
    assert "un-pivoted" not in out


# ---- discovery: .docx is now an accepted source suffix ----

def test_docx_is_a_discovered_suffix():
    src = (pathlib.Path(__file__).resolve().parents[1] / "vaxtract" / "extraction_agent.py").read_text()
    assert '".pdf", ".xlsx", ".docx"' in src
