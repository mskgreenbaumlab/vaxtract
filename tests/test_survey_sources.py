"""#3 source-discovery: survey_sources inventories every supplement in one call so the agent can
LOCATE a manifest hidden among many oddly-named files (the 39972124 failure mode)."""
import pathlib
import openpyxl
import docx as _docx
import agent_core


def _xlsx(path, sheets):
    wb = openpyxl.Workbook(); wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(name)
        for r in rows:
            ws.append(r)
    wb.save(path)


def test_survey_surfaces_hidden_peptide_table_by_its_columns(tmp_path):
    # a peptide manifest hidden among figure-source-data sheets/files (39972124 shape)
    _xlsx(tmp_path / "MOESM4.xlsx", {"Figure 2a": [["x", "y"], [1, 2]],
                                     "list_of_diff_genes": [["gene"], ["TP53"]]})
    _xlsx(tmp_path / "MOESM7.xlsx", {"treated_neoantigens":
                                     [["Patient", "Peptide", "HLA"], ["N06", "DDLRYK", "A*11:01"]]})
    out = agent_core.survey_sources(str(tmp_path))
    assert "MOESM7.xlsx" in out and "treated_neoantigens" in out
    assert "Patient" in out and "Peptide" in out and "HLA" in out      # found by its columns
    assert "Figure 2a" in out                                          # the noise is inventoried too


def test_survey_docx_caption_and_table_header(tmp_path):
    d = _docx.Document()
    d.add_paragraph("Supplementary Table S2. Treated neoantigen peptides")
    t = d.add_table(rows=1, cols=3)
    for c, h in zip(t.rows[0].cells, ["Patient", "Peptide", "MHC"]):
        c.text = h
    d.save(tmp_path / "S2.docx")
    out = agent_core.survey_sources(str(tmp_path))
    assert "S2.docx" in out and "Treated neoantigen peptides" in out
    assert "Patient" in out and "MHC" in out


def test_survey_single_file_and_dims(tmp_path):
    p = tmp_path / "one.xlsx"
    _xlsx(p, {"Sheet1": [["A", "B"], [1, 2], [3, 4]]})
    out = agent_core.survey_sources(str(p))
    assert "one.xlsx" in out and "Sheet1" in out and "3x2" in out      # rows x cols


def test_survey_recurses_subdirs(tmp_path):
    sub = tmp_path / "supps"; sub.mkdir()
    _xlsx(sub / "deep.xlsx", {"S": [["h"], ["v"]]})
    assert "deep.xlsx" in agent_core.survey_sources(str(tmp_path))


def test_survey_ignores_lockfiles_and_reports_empty(tmp_path):
    (tmp_path / "~$tmp.xlsx").write_text("lock")
    assert "no .xlsx/.pdf/.docx" in agent_core.survey_sources(str(tmp_path))
    assert "path not found" in agent_core.survey_sources(str(tmp_path / "nope"))


def test_survey_byte_cap_lists_skipped(tmp_path):
    for i in range(6):
        _xlsx(tmp_path / f"f{i}.xlsx", {"S": [["header_col_alpha", "header_col_beta"], [1, 2]]})
    out = agent_core.survey_sources(str(tmp_path), max_chars=200)   # tiny cap -> must skip some
    assert "capped:" in out and "not shown" in out


def test_survey_tool_is_registered():
    src = (pathlib.Path(__file__).resolve().parents[1] / "vaxtract" / "extraction_agent.py").read_text()
    assert '@tool("survey_sources"' in src
    assert "survey_sources," in src                       # in the server tools=[...] list
    assert "mcp__antvac__survey_sources" in src           # in the allowed-tools list


def test_survey_flags_per_patient_sheet_family(tmp_path):
    # step-3 fix (2026-06-09): >=5 same-schema tabs -> flag a family + prescribe ONE bulk add_table
    # with a copy-ready sheets=[...] list (the 33064988 ~34 IAP-tab under-sweep, confirmed live).
    sheets = {f"IAP-P{i}": [["Peptide", "SFC", "Response"], [f"PEP{i}", 120, "yes"]] for i in range(6)}
    _xlsx(tmp_path / "MOESM6.xlsx", sheets)
    out = agent_core.survey_sources(str(tmp_path))
    assert "PER-PATIENT SHEET FAMILY" in out and "6 sheets share one schema" in out
    assert "sheets=[" in out and "IAP-P0" in out and "IAP-P5" in out   # copy-ready list
    assert "__sheet__" in out

def test_survey_no_family_flag_below_threshold(tmp_path):
    # only 3 same-schema tabs -> single reads are fine, no family directive
    sheets = {f"S{i}": [["Peptide", "SFC"], [f"P{i}", 1]] for i in range(3)}
    _xlsx(tmp_path / "MOESM9.xlsx", sheets)
    out = agent_core.survey_sources(str(tmp_path))
    assert "PER-PATIENT SHEET FAMILY" not in out
